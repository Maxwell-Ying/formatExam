import random
from typing import List

from docx.text.paragraph import Paragraph

space = "       "
bracket = "（    ）"


def draw_one_multi_pron(paragraph: Paragraph, letter: str, bold=False):
    run1 = paragraph.add_run(space)
    run1.bold = bold
    run1.font.underline = True
    paragraph.add_run('\n')
    run2 = paragraph.add_run(letter)
    run2.bold = bold
    paragraph.add_run('\n')
    run3 = paragraph.add_run(bracket)
    run3.bold = bold
    paragraph.add_run('\n')


def draw_two_multi_pron(paragraph: Paragraph, letter1: str, letter2: str, bold=False):
    draw_empty(paragraph)

    run6 = paragraph.add_run(letter1)
    run6.bold = bold
    paragraph.add_run(space*4)
    run8 = paragraph.add_run(letter2)
    run8.bold = bold
    paragraph.add_run('\n')

    draw_empty(paragraph)


def draw_empty(paragraph: Paragraph):
    paragraph.add_run(space)
    run1 = paragraph.add_run(space)
    run1.underline = True
    paragraph.add_run(bracket)
    paragraph.add_run(space + space)
    run4 = paragraph.add_run(space)
    run4.underline = True
    paragraph.add_run(bracket)
    paragraph.add_run('\n')


def draw_like(paragraph: Paragraph, words: List[str], bold=False):
    words = random.sample(words, 2)
    for word in words:
        run = paragraph.add_run(word)
        run.bold = bold
        run2 = paragraph.add_run(bracket*3)
        run2.bold = bold
        paragraph.add_run(space)


def draw_new_line(paragraph: Paragraph):
    paragraph.add_run("\n")


def draw_letter(paragraph: Paragraph, word: str, bold=False):
    run = paragraph.add_run(word)
    run.bold = bold


def draw_letter_space(paragraph: Paragraph, letter: str, bold=False):
    run = paragraph.add_run(letter + bracket * 3)
    run.bold = bold
