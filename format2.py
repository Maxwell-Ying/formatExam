import os

import docx
from docx.text.paragraph import Paragraph

from common_format import format_normal, format_heading, format_heading_run, format_paragraph_title, format_paragraph
from draw import draw_bold, draw_like, draw_two_multi_pron, draw_letter, draw_new_line
from get_multi_pronouce import load_multi_pronounce
from load_new_letters import load_new_letters
from load_new_words import load_new_words
from load_similar_letters import load_similar_letters

day_in_week = "零一二三四五六"
overall_font = u"宋体"
brackets = "（   ）  （   ）  （   ）"
level = 2

all_words = load_new_words(level)
all_letters = load_new_letters(level)
all_similarities = load_similar_letters(level)
all_pronounce = load_multi_pronounce(level)

for index in all_words.keys():
    if index not in all_letters:
        continue
    doc_name = day_in_week[2] + "年级第" + str(index) + "周课内巩固"
    document = docx.Document()
    format_normal(document)

    heading: Paragraph = document.add_heading('', level=0)
    heading_format = heading.paragraph_format
    format_heading(heading)
    run_head = heading.add_run(doc_name)
    format_heading_run(run_head)

    paragraph1 = document.add_paragraph()
    format_paragraph_title(paragraph1, 1)
    format_paragraph(paragraph1)

    letter_count = 0
    line_count = 0
    for word in all_words[index]:
        if len(word) > 4:
            continue
        if letter_count + len(word) > 16:
            if line_count > 0:
                paragraph1.add_run("\n")
                break
            else:
                paragraph1.add_run("\n\n")
                letter_count = 0
            line_count = len(word)
        else:
            letter_count += len(word)
        draw_bold(paragraph1, word + " ")
    draw_new_line(paragraph1)

    paragraph2 = document.add_paragraph()
    format_paragraph_title(paragraph2, 2)
    format_paragraph(paragraph2)

    selected_letters = all_letters[index]
    for x in range(min(6, len(selected_letters))):
        draw_letter(paragraph2, selected_letters[x])
        if x % 2 == 1 and x != 5:
            draw_new_line(paragraph2)

    paragraph3 = document.add_paragraph()
    format_paragraph_title(paragraph3, 3)
    format_paragraph(paragraph3)

    if len(all_similarities) > 0:
        candidate = all_similarities[0]
        all_similarities = all_similarities[1:]
        draw_like(paragraph3, candidate[:2])
        draw_like(paragraph3, candidate[2:])
        draw_new_line(paragraph3)

    paragraph4 = document.add_paragraph()
    format_paragraph_title(paragraph4, 4)
    format_paragraph(paragraph4)

    if len(all_pronounce) > 1:
        draw_two_multi_pron(paragraph4, all_pronounce[0], all_pronounce[1])
        all_pronounce = all_pronounce[2:]

    document.save(os.path.join("result_"+str(level), str(level)+"-"+str(index) + ".docx"))
    # break
