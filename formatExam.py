import os
import sys
import random

import docx
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
import divide
from docx.shared import Length

day_in_week = "一二三四五六"


def get_words(fin):
    ret = []
    for line in fin.readlines():
        line = line.strip()
        word = line.split("\t")[0]
        ret.append(word)
    return ret

#
# words = []
#
# for file in os.listdir("data"):
#     with open(os.path.join("data", file)) as f:
#         words += get_words(f)
#
# words = [i for i in words if 2 <= len(i) <= 4]


for i in range(1):
    degree = i+1
    for j in range(1):
        ws, new_word = divide.get_words()
        doc_name = day_in_week[i] + "年级第" + str(j+1) + "周课内巩固"
        document = docx.Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_heading(doc_name, level=2)
        paragraph = document.add_paragraph("给下列词语注音并抄写\n")
        paragraph_format = paragraph.paragraph_format
        paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.5
        for _ in range(10):
            paragraph.add_run(random.choice(ws) + " ").bold = True
        paragraph.add_run("\n")

        paragraph2 = document.add_paragraph("生字组词\n")
        paragraph_format2 = paragraph.paragraph_format
        paragraph2.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format2.line_spacing = 1.5
        for x in range(6):
            run = paragraph2.add_run(random.choice(new_word) + "(   )  (   )  (   ) ")
            run.font.name = "黑体"
            run.bold = True
            if x % 2 == 1:
                paragraph2.add_run("\n")

        document.save(str(degree) + "-" + str(j+1) + ".docx")
        



