import os
import sys
import random

import docx
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
# import divide
from docx.shared import Length, Pt

from draw import draw_one_multi_pron, draw_two_multi_pron, draw_like
from load_new_words import load_new_words
from load_book import load_pages
from load_likes import load_likes
from get_multi_pronouce import get_multi_pronounce
from load_likes import get_similar_letters

day_in_week = "一二三四五六"
overall_font = u"宋体"
brackets = "（   ）  （   ）  （   ）"

all_words = load_new_words(6)
all_letters = load_pages()
# all_likes = load_likes()

for index in all_words.keys():
    if index not in all_letters:
        continue
    doc_name = "六" + "年级第" + str(index) + "周课内巩固"
    document = docx.Document()
    document.styles['Normal'].font.name = overall_font
    document.styles['Normal'].font.size = Pt(15)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    heading = document.add_heading('', level=1)
    heading_format = heading.paragraph_format
    heading_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_head = heading.add_run(doc_name)
    run_head.font.name = overall_font
    run_head.font.size = Pt(17)
    run_head._element.rPr.rFonts.set(qn('w:eastAsia'), overall_font)
    run_head.bold = True

    paragraph = document.add_paragraph()
    head1 = paragraph.add_run("一、给下列词语注音并抄写\n")
    head1.font.name = overall_font
    head1.font.size = Pt(15)
    head1._element.rPr.rFonts.set(qn('w:westAsia'), overall_font)
    head1.bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.5
    letter_count = 0
    line_count = 0
    for word in all_words[index]:
        if len(word) > 4:
            continue
        if letter_count + len(word) > 16:
            paragraph.add_run("\n\n").font.size = Pt(15)
            letter_count = len(word)
            line_count += 1
            if line_count >= 2:
                break
        else:
            letter_count += len(word)
        run = paragraph.add_run(word + " ")
        run.bold = True
        run.font.size = Pt(15)

    paragraph2 = document.add_paragraph()
    head2 = paragraph2.add_run("二、生字组词\n")
    head2.font.size = Pt(15)
    head2.bold = True
    paragraph_format2 = paragraph2.paragraph_format
    paragraph2.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format2.line_spacing = 1.5

    for x in range(min(6, len(all_letters[index]))):
        run = paragraph2.add_run(all_letters[index][x] + brackets + "   ")
        run.font.size = Pt(15)
        if x % 2 == 1:
            paragraph2.add_run("\n")

    paragraph3 = document.add_paragraph()
    head3 = paragraph3.add_run("三、形近字辨析\n")
    head3.font.size = Pt(15)
    head3.bold = True
    paragraph_format3 = paragraph3.paragraph_format
    paragraph_format3.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # like1, like2 = all_likes[:2]
    # for x in range(len(like1)):
    #     run = paragraph3.add_run(like1[x] + brackets + "   " + like2[x] + brackets + "\n")
    #     run.bold = True
    candidate = get_similar_letters(all_words[index])
    if len(candidate) < 2:
        print("two few letter satisfy like condition")
    for like in random.sample(candidate, min(len(candidate), 3)):
        draw_like(paragraph3, like)

    paragraph4 = document.add_paragraph()
    head4 = paragraph4.add_run("四、多音字辨析\n")
    head4.font.name = overall_font
    head4.font.size = Pt(15)
    head4._element.rPr.rFonts.set(qn('w:westAsia'), overall_font)
    head4.bold = True
    paragraph_format4 = paragraph4.paragraph_format
    paragraph_format4.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    multi_pronounce = get_multi_pronounce(all_words[index])
    if len(multi_pronounce) == 0:
        print("第"+str(index)+"周生字中没有足够的多音字")
    elif len(multi_pronounce) == 1:
        draw_one_multi_pron(paragraph4, multi_pronounce[0])
    else:
        draw_two_multi_pron(paragraph4, multi_pronounce[0], multi_pronounce[1])
        if len(multi_pronounce) >= 4:
            draw_two_multi_pron(paragraph4, multi_pronounce[2], multi_pronounce[3])

    document.save(os.path.join("result", "6" + "-" + str(index) + ".docx"))
    # break
